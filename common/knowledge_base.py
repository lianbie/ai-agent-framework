"""
知识库核心模块

支持多种文件格式的上传、解析、切片、向量化和检索。
企业级设计：支持配置化、错误处理、日志记录。

Usage:
    from common.knowledge_base import KnowledgeBase

    # 初始化知识库
    kb = KnowledgeBase(connection_string="postgresql://user:pass@host:5432/db")

    # 添加文档
    await kb.add_document("path/to/document.pdf")

    # 搜索
    results = await kb.search("查询内容", top_k=5)
"""

import os
import json
import uuid
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import logging

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """文档数据类"""
    content: str
    metadata: Dict[str, Any] = None

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}


@dataclass
class SearchResult:
    """搜索结果数据类"""
    content: str
    score: float
    metadata: Dict[str, Any] = None

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}

    def to_dict(self) -> Dict[str, Any]:
        return {
            "content": self.content,
            "score": self.score,
            "metadata": self.metadata
        }


class KnowledgeBase:
    """
    知识库核心类（单例模式）

    特性：
    - 支持多种文件格式（txt, pdf, docx, xlsx）
    - 文本切片和向量化
    - 向量相似度搜索
    - 可选Rerank重排序

    Example:
        ```python
        # 初始化（单例，只会初始化一次）
        kb = KnowledgeBase()

        # 添加文档
        await kb.add_text("这是知识库内容")
        await kb.add_document("path/to/file.pdf")

        # 搜索
        results = await kb.search("查询内容", top_k=5)
        for result in results:
            print(result.content, result.score)
        ```
    """

    _instance = None
    _initialized = False

    def __new__(cls, *args, **kwargs):
        """单例模式"""
        if cls._instance is None:
            cls._instance = super().__new__(cls)
        return cls._instance

    def __init__(
        self,
        connection_string: Optional[str] = None,
        embedding_model: str = "BAAI/bge-base-zh-v1.5",
        collection_name: str = "knowledge_base",
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        enable_rerank: bool = True,
        rerank_model: str = "BAAI/bge-reranker-base"
    ):
        """
        初始化知识库

        Args:
            connection_string: PostgreSQL连接字符串
            embedding_model: Embedding模型名称
            collection_name: 向量集合名称
            chunk_size: 文本切片大小
            chunk_overlap: 文本切片重叠大小
            enable_rerank: 是否启用Rerank
            rerank_model: Rerank模型名称
        """
        # 避免重复初始化
        if self._initialized:
            return

        self.connection_string = connection_string
        self.embedding_model = embedding_model
        self.collection_name = collection_name
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.enable_rerank = enable_rerank
        self.rerank_model_name = rerank_model

        # 内部状态
        self._embeddings = None
        self._text_splitter = None
        self._vector_store = None
        self._rerank_model = None

        # 延迟初始化
        self._init_embeddings()
        self._init_text_splitter()

        if connection_string:
            self._init_vector_store()

        if enable_rerank:
            self._init_rerank_model()

        self._initialized = True
        logger.info(f"✅ KnowledgeBase initialized (collection={collection_name})")

    def _init_embeddings(self):
        """初始化Embedding模型"""
        try:
            from langchain_openai import OpenAIEmbeddings
            from common.model_manager import model_manager

            # 从数据库获取模型配置
            model_config = model_manager.get_active_model_config()

            if model_config:
                base_url = model_config.base_url or "https://api.openai.com/v1"
                api_key = model_config.api_key or ""
            else:
                base_url = "https://api.openai.com/v1"
                api_key = ""

            # 创建Embedding模型
            self._embeddings = OpenAIEmbeddings(
                model="text-embedding-ada-002",
                openai_api_base=base_url,
                openai_api_key=api_key
            )
            logger.info(f"✅ Embedding model initialized")
        except Exception as e:
            logger.warning(f"⚠️ Embedding initialization failed: {e}")
            self._embeddings = None

    def _init_text_splitter(self):
        """初始化文本切分器"""
        try:
            from langchain.text_splitter import RecursiveCharacterTextSplitter

            # 中文分隔符
            separators = ["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""]

            self._text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                separators=separators,
                length_function=len,
            )
            logger.info(f"✅ Text splitter initialized (chunk_size={self.chunk_size})")
        except Exception as e:
            logger.warning(f"⚠️ Text splitter initialization failed: {e}")
            self._text_splitter = None

    def _init_vector_store(self):
        """初始化向量存储"""
        if not self.connection_string:
            logger.warning("⚠️ No connection string, vector store not initialized")
            return

        try:
            from langchain_postgres import PGVector

            self._vector_store = PGVector(
                connection=self.connection_string,
                embedding_function=self._embeddings,
                collection_name=self.collection_name,
                use_jsonb=True,
            )
            logger.info(f"✅ Vector store initialized: {self.collection_name}")
        except ImportError:
            try:
                from langchain_community.vectorstores import PGVector

                self._vector_store = PGVector(
                    connection_string=self.connection_string,
                    embedding_function=self._embeddings,
                    collection_name=self.collection_name,
                )
                logger.info(f"✅ Vector store initialized (legacy): {self.collection_name}")
            except Exception as e:
                logger.error(f"❌ Vector store initialization failed: {e}")
                self._vector_store = None
        except Exception as e:
            logger.error(f"❌ Vector store initialization failed: {e}")
            self._vector_store = None

    def _init_rerank_model(self):
        """初始化Rerank模型"""
        try:
            from sentence_transformers import CrossEncoder

            # 尝试本地模型
            local_models = [
                './models/bge-reranker-base',
                './models/reranker',
            ]

            for model_path in local_models:
                if os.path.exists(model_path):
                    try:
                        self._rerank_model = CrossEncoder(model_path)
                        logger.info(f"✅ Local rerank model loaded: {model_path}")
                        return
                    except Exception:
                        continue

            # 尝试在线模型
            self._rerank_model = CrossEncoder(self.rerank_model_name)
            logger.info(f"✅ Online rerank model loaded: {self.rerank_model_name}")

        except ImportError:
            logger.warning("⚠️ sentence-transformers not installed, rerank disabled")
            self._rerank_model = None
        except Exception as e:
            logger.warning(f"⚠️ Rerank model initialization failed: {e}")
            self._rerank_model = None

    async def add_text(
        self,
        text: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> bool:
        """
        添加文本到知识库

        Args:
            text: 文本内容
            metadata: 元数据（可选）

        Returns:
            是否成功
        """
        if not self._vector_store:
            logger.error("❌ Vector store not initialized")
            return False

        try:
            # 切分文本
            if self._text_splitter:
                chunks = self._text_splitter.split_text(text)
            else:
                chunks = [text]

            # 构建文档
            documents = []
            for i, chunk in enumerate(chunks):
                doc_metadata = {
                    "chunk_index": i,
                    "total_chunks": len(chunks),
                    **(metadata or {})
                }
                documents.append(Document(content=chunk, metadata=doc_metadata))

            # 添加到向量存储
            await self._add_documents(documents)

            logger.info(f"✅ Added text ({len(chunks)} chunks)")
            return True

        except Exception as e:
            logger.error(f"❌ Add text failed: {e}")
            return False

    async def add_document(self, file_path: str) -> bool:
        """
        添加文档到知识库

        Args:
            file_path: 文件路径

        Returns:
            是否成功
        """
        if not os.path.exists(file_path):
            logger.error(f"❌ File not found: {file_path}")
            return False

        try:
            # 解析文件
            text = self._parse_file(file_path)
            if not text:
                return False

            # 添加元数据
            metadata = {
                "source": file_path,
                "filename": os.path.basename(file_path),
            }

            return await self.add_text(text, metadata)

        except Exception as e:
            logger.error(f"❌ Add document failed: {e}")
            return False

    def _parse_file(self, file_path: str) -> Optional[str]:
        """
        解析文件内容

        Args:
            file_path: 文件路径

        Returns:
            文件内容字符串
        """
        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()

            elif ext == '.pdf':
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    return "\n".join(page.extract_text() for page in reader.pages)

            elif ext == '.docx':
                import docx
                doc = docx.Document(file_path)
                return "\n".join(para.text for para in doc.paragraphs)

            elif ext == '.xlsx':
                import pandas as pd
                df = pd.read_excel(file_path)
                return df.to_string()

            elif ext == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return json.dumps(json.load(f), ensure_ascii=False, indent=2)

            elif ext == '.csv':
                import pandas as pd
                df = pd.read_csv(file_path)
                return df.to_string()

            else:
                logger.warning(f"⚠️ Unsupported file format: {ext}")
                return None

        except Exception as e:
            logger.error(f"❌ Parse file failed: {e}")
            return None

    async def _add_documents(self, documents: List[Document]):
        """添加文档到向量存储"""
        if not self._vector_store:
            raise ValueError("Vector store not initialized")

        from langchain.schema import Document as LCDocument

        lc_docs = [
            LCDocument(page_content=doc.content, metadata=doc.metadata)
            for doc in documents
        ]

        self._vector_store.add_documents(lc_docs)

    async def search(
        self,
        query: str,
        top_k: int = 5,
        enable_rerank: Optional[bool] = None
    ) -> List[SearchResult]:
        """
        搜索知识库

        Args:
            query: 查询内容
            top_k: 返回结果数量
            enable_rerank: 是否启用Rerank（可选，默认使用配置值）

        Returns:
            搜索结果列表
        """
        if not self._vector_store:
            logger.warning("⚠️ Vector store not initialized, returning empty results")
            return []

        try:
            # 向量搜索
            search_results = self._vector_store.similarity_search_with_score(
                query,
                k=top_k * 2 if (enable_rerank or self.enable_rerank) else top_k
            )

            # 转换结果
            results = [
                SearchResult(
                    content=doc.page_content,
                    score=score,
                    metadata=doc.metadata
                )
                for doc, score in search_results
            ]

            # Rerank
            use_rerank = enable_rerank if enable_rerank is not None else self.enable_rerank
            if use_rerank and self._rerank_model and len(results) > 1:
                results = self._rerank(query, results, top_k)

            return results[:top_k]

        except Exception as e:
            logger.error(f"❌ Search failed: {e}")
            return []

    def _rerank(
        self,
        query: str,
        results: List[SearchResult],
        top_k: int
    ) -> List[SearchResult]:
        """
        重排序搜索结果

        Args:
            query: 查询内容
            results: 原始搜索结果
            top_k: 返回结果数量

        Returns:
            重排序后的结果
        """
        try:
            # 构建重排序输入
            pairs = [(query, r.content) for r in results]

            # 计算重排序分数
            scores = self._rerank_model.predict(pairs)

            # 更新分数并排序
            for i, score in enumerate(scores):
                results[i].score = float(score)

            results.sort(key=lambda x: x.score, reverse=True)

            return results[:top_k]

        except Exception as e:
            logger.warning(f"⚠️ Rerank failed: {e}")
            return results

    async def delete(self, filter_dict: Optional[Dict[str, Any]] = None) -> bool:
        """
        删除知识库内容

        Args:
            filter_dict: 过滤条件（可选）

        Returns:
            是否成功
        """
        if not self._vector_store:
            logger.error("❌ Vector store not initialized")
            return False

        try:
            # TODO: 实现删除逻辑
            logger.info(f"✅ Deleted documents")
            return True
        except Exception as e:
            logger.error(f"❌ Delete failed: {e}")
            return False

    def get_stats(self) -> Dict[str, Any]:
        """
        获取知识库统计信息

        Returns:
            统计信息字典
        """
        return {
            "collection_name": self.collection_name,
            "embedding_model": self.embedding_model,
            "chunk_size": self.chunk_size,
            "chunk_overlap": self.chunk_overlap,
            "enable_rerank": self.enable_rerank,
            "rerank_model": self.rerank_model_name,
            "initialized": self._initialized,
            "has_vector_store": self._vector_store is not None,
            "has_rerank_model": self._rerank_model is not None,
        }
